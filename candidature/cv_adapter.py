"""
cv_adapter.py — Génération de CV adaptés par offre
===================================================

CONCEPT CLÉ : personnalisation dynamique.

On ne change pas le CONTENU du CV (pas de mensonge),
on change l'ORDRE et la MISE EN AVANT des compétences
selon ce que l'offre recherche.

Exemple :
- Offre "Développeur Java" → Java monte en premier dans les langages
- Offre "Data Analyst" → Python + SQL montent, React descend
- Offre "DevOps" → Linux + Git montent, on met en avant le projet réseau

L'IA analyse l'offre et décide de l'ordre optimal.
"""

import json
import re
import unicodedata
from pathlib import Path
from typing import Optional

try:
    import anthropic
except ImportError:
    raise ImportError("pip install anthropic")

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("pip install python-docx")

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

try:
    from logger import get_logger
    log = get_logger("candidature.cv_adapter")
except ImportError:
    import logging
    log = logging.getLogger(__name__)


PROFIL_PATH = Path(__file__).parent.parent / "qualification" / "profil.yaml"


def charger_profil_brut(chemin: Path = PROFIL_PATH) -> str:
    """Charge le profil en texte brut pour injection dans les prompts."""
    if not chemin.exists():
        raise FileNotFoundError(f"❌ Profil introuvable : {chemin}")
    return chemin.read_text(encoding="utf-8")


def charger_profil_dict(chemin: Path = PROFIL_PATH) -> dict:
    """
    Charge le profil en dict structuré pour extraire les infos personnelles.
    
    Utilise yaml si disponible, sinon un parsing simplifié clé: valeur.
    """
    texte = charger_profil_brut(chemin)
    if HAS_YAML:
        return yaml.safe_load(texte) or {}
    # Fallback minimaliste : ne couvre pas les listes imbriquées
    profil = {}
    for ligne in texte.splitlines():
        if ":" in ligne and not ligne.strip().startswith("#"):
            cle, _, valeur = ligne.partition(":")
            profil[cle.strip()] = valeur.strip()
    return profil


def nom_fichier_safe(texte: str, max_len: int = 60) -> str:
    """
    Transforme un texte en nom de fichier safe (ASCII, sans accents ni spéciaux).
    
    'Développeur Full-Stack / Énergie' → 'Developpeur_Full-Stack_Energie'
    """
    # Normaliser les accents (é → e, ç → c, etc.)
    texte = unicodedata.normalize("NFKD", texte)
    texte = texte.encode("ascii", "ignore").decode("ascii")
    # Remplacer espaces et slashs
    texte = texte.replace(" ", "_").replace("/", "-")
    # Garder uniquement alphanum, tiret, underscore
    texte = re.sub(r"[^a-zA-Z0-9_-]", "", texte)
    return texte[:max_len]


class CVAdapter:
    """Génère des CV adaptés à chaque offre."""

    def __init__(
        self,
        api_key: str,
        model: str = "claude-sonnet-4-20250514",
        timeout: float = 30.0,
    ):
        self.client = anthropic.Anthropic(api_key=api_key, timeout=timeout)
        self.model = model
        self.profil_brut = charger_profil_brut()
        self.profil_dict = charger_profil_dict()
        log.info(f"CVAdapter initialisé (modèle: {model})")

    # ------------------------------------------------------------------
    # Extraction des infos personnelles depuis profil.yaml
    # ------------------------------------------------------------------

    def _info(self, *cles: str, defaut: str = "") -> str:
        """Cherche une valeur dans le profil dict par clé(s) possible(s)."""
        for cle in cles:
            val = self.profil_dict.get(cle)
            if val:
                return str(val)
        return defaut

    def _liste(self, *cles: str, defaut: list | None = None) -> list:
        """Cherche une liste dans le profil dict."""
        for cle in cles:
            val = self.profil_dict.get(cle)
            if isinstance(val, list):
                return val
        return defaut or []

    @property
    def nom(self) -> str:
        return self._info("nom", defaut="Candidat")

    @property
    def email(self) -> str:
        return self._info("email", "mail", defaut="")

    @property
    def telephone(self) -> str:
        return self._info("telephone", "tel", "téléphone", defaut="")

    @property
    def localisation(self) -> str:
        return self._info("localisation", "lieu", "region", defaut="")

    @property
    def linkedin(self) -> str:
        return self._info("linkedin", defaut="")

    @property
    def github(self) -> str:
        return self._info("github", defaut="")

    # ------------------------------------------------------------------
    # Analyse IA de l'offre
    # ------------------------------------------------------------------

    def analyser_offre(self, offre: dict) -> Optional[dict]:
        """
        Demande à Claude d'analyser l'offre et de réordonner les compétences.
        Retourne un dict avec l'ordre optimal des sections.
        """
        prompt = f"""Tu es un expert en recrutement tech. Analyse cette offre et le profil du candidat.

PROFIL :
{self.profil_brut}

OFFRE :
- Titre : {offre.get('titre', '')}
- Entreprise : {offre.get('entreprise', '')}
- Description : {offre.get('description', '')}

Réordonne les compétences et expériences du candidat pour MAXIMISER le match avec cette offre.
Ne change PAS le contenu, change juste l'ORDRE de présentation.

Réponds UNIQUEMENT en JSON :
{{
    "titre_cv": "Développeur [adapté au poste]",
    "accroche": "2 phrases max résumant pourquoi ce candidat est pertinent pour CE poste",
    "langages_ordonnes": ["langage le plus pertinent en premier", "..."],
    "frameworks_ordonnes": ["framework le plus pertinent en premier", "..."],
    "bdd_ordonnees": ["BDD la plus pertinente en premier", "..."],
    "outils_ordonnes": ["outil le plus pertinent en premier", "..."],
    "projet_star_titre": "le TITRE EXACT du projet le plus pertinent (ex: Agent IA de recherche d'alternance)",
    "competences_ia_pertinentes": ["compétence IA/auto pertinente", "..."],
    "mots_cles_offre": ["3 mots-clés de l'offre à mettre en avant"]
}}"""

        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=600,  # 500 était trop juste, risque de JSON tronqué
                messages=[{"role": "user", "content": prompt}],
            )
            texte = response.content[0].text
            try:
                return json.loads(texte)
            except json.JSONDecodeError:
                debut = texte.index("{")
                fin = texte.rindex("}") + 1
                return json.loads(texte[debut:fin])

        except anthropic.APITimeoutError:
            log.error(f"Timeout API pour analyse CV : {offre.get('titre', '?')}")
            print(f"  ⏱️  Timeout API — CV par défaut")
            return None

        except Exception as e:
            log.error(f"Erreur analyse CV : {e}")
            print(f"  ⚠️  Erreur analyse : {e}")
            return None

    # ------------------------------------------------------------------
    # Génération du document Word
    # ------------------------------------------------------------------

    def generer_cv(self, offre: dict, output_dir: Path) -> Optional[Path]:
        """
        Génère un CV Word adapté à l'offre.
        Les infos personnelles viennent de profil.yaml, pas du code.
        """
        log.info(f"Adaptation du CV pour : {offre.get('titre', '?')} @ {offre.get('entreprise', '?')}")
        print(f"  📄 Adaptation du CV pour : {offre.get('titre', '?')}")

        # --- Analyse IA ---
        analyse = self.analyser_offre(offre)
        if not analyse:
            log.warning("Analyse IA échouée, utilisation des valeurs par défaut")
            print(f"  ⚠️  Analyse échouée, CV par défaut")
            analyse = {
                "titre_cv": "Développeur logiciel",
                "accroche": "Développeur polyvalent en recherche d'alternance.",
                "langages_ordonnes": ["Python", "JavaScript", "Java", "HTML/CSS", "C"],
                "frameworks_ordonnes": ["React", "Node.js"],
                "bdd_ordonnees": ["PostgreSQL", "MySQL", "MongoDB"],
                "outils_ordonnes": ["Linux", "VS Code", "Git", "VirtualBox", "Wireshark"],
                "competences_ia_pertinentes": ["Web scraping", "API Claude"],
                "projet_star_titre": "",
                "mots_cles_offre": [],
            }

        # --- Créer le document ---
        doc = Document()

        # Marges
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Style par défaut
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)

        # === EN-TÊTE (depuis profil.yaml) ===
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.nom.upper())
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(30, 60, 90)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(analyse.get("titre_cv", "Développeur logiciel"))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)

        # Contact — construit dynamiquement depuis le profil
        contact_parts = [p for p in [self.email, self.telephone, self.localisation] if p]
        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(contact_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

        liens_parts = []
        if self.linkedin:
            liens_parts.append(f"LinkedIn : {self.linkedin}")
        if self.github:
            liens_parts.append(f"GitHub : {self.github}")
        if liens_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(liens_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Ligne séparatrice
        self._ajouter_separateur(doc)

        # === ACCROCHE ===
        accroche = analyse.get("accroche", "")
        if accroche:
            p = doc.add_paragraph()
            run = p.add_run(accroche)
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(60, 60, 60)
            p.paragraph_format.space_after = Pt(10)

        # === COMPÉTENCES TECHNIQUES ===
        self._titre_section(doc, "Compétences techniques")

        competences = [
            ("Langages", " — ".join(analyse.get("langages_ordonnes", ["Python", "JavaScript", "Java", "HTML/CSS", "C"]))),
            ("Frameworks", " — ".join(analyse.get("frameworks_ordonnes", ["React", "Node.js"]))),
            ("Bases de données", " — ".join(analyse.get("bdd_ordonnees", ["PostgreSQL", "MySQL", "MongoDB"]))),
            ("Outils", " — ".join(analyse.get("outils_ordonnes", ["Linux", "VS Code", "Git"]))),
            ("Méthodes", "Agile — Architecture MVC — UML"),
        ]

        ia_comp = analyse.get("competences_ia_pertinentes", [])
        if ia_comp:
            competences.append(("IA & Automatisation", " — ".join(ia_comp)))

        for label, valeur in competences:
            p = doc.add_paragraph()
            run = p.add_run(f"{label} : ")
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(valeur)
            run.font.size = Pt(10)

        # === EXPÉRIENCE PROFESSIONNELLE ===
        self._titre_section(doc, "Expérience professionnelle")

        # TODO: Externaliser les expériences dans profil.yaml
        # Pour l'instant, hardcodé (identique à l'original)
        p = doc.add_paragraph()
        run = p.add_run("Développeur Full Stack — Bizlab")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run("  |  Mai 2023 — Juillet 2023")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

        missions = [
            "Analyse des besoins clients et définition des exigences techniques et fonctionnelles",
            "Développement d'un site web dynamique et responsive (PHP, HTML, CSS, JavaScript)",
            "Gestion de la base de données PostgreSQL",
        ]
        for m in missions:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(m)
            run.font.size = Pt(9.5)

        # === PROJETS ===
        self._titre_section(doc, "Projets")

        # TODO: Externaliser les projets dans profil.yaml
        projets = [
            {
                "titre": "Agent IA de recherche d'alternance",
                "tech": "Python, API Claude, Web Scraping, BeautifulSoup, Algolia API",
                "desc": "Pipeline automatisé : collecte d'offres (WTTJ), scoring IA, génération de candidatures personnalisées, dashboard de suivi. Reverse engineering d'API, prompt engineering, génération de documents Word/PDF."
            },
            {
                "titre": "API REST + Application React",
                "tech": "Node.js, React, REST API",
                "desc": "Application web de gestion de stock avec API REST en Node.js et interface React."
            },
            {
                "titre": "Conception d'un réseau (L3)",
                "tech": "VirtualBox, Wireshark, DNS, DHCP, HTTPS",
                "desc": "Adressage et routage IP, mise en place de services réseau complets."
            },
            {
                "titre": "Jeu Luzhanqi en Java (L2)",
                "tech": "Java, JavaSwing, Git, UML",
                "desc": "Conception MVC, interface graphique, gestion des événements utilisateur."
            },
        ]

        # Réordonner : mettre le projet star en premier (matching amélioré)
        projet_star_titre = analyse.get("projet_star_titre", "")
        if projet_star_titre:
            self._remonter_projet_star(projets, projet_star_titre)

        for proj in projets:
            p = doc.add_paragraph()
            run = p.add_run(proj["titre"])
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(f"  |  {proj['tech']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

            p = doc.add_paragraph()
            run = p.add_run(proj["desc"])
            run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(6)

        # === FORMATION ===
        self._titre_section(doc, "Formation")

        # TODO: Externaliser dans profil.yaml
        p = doc.add_paragraph()
        run = p.add_run("Licence Informatique — CY Université")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run("  |  2019 — 2024")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # === LANGUES & INTÉRÊTS ===
        self._titre_section(doc, "Langues et centres d'intérêt")

        langues = self._liste("langues", defaut=[])
        if langues:
            p = doc.add_paragraph()
            run = p.add_run("Langues : ")
            run.bold = True
            run = p.add_run(", ".join(langues) if isinstance(langues[0], str) else
                           ", ".join(f"{l}" for l in langues))
        else:
            p = doc.add_paragraph()
            run = p.add_run("Langues : ")
            run.bold = True
            run = p.add_run("Français (natif), Anglais (avancé)")

        interets = self._liste("interets", "intérêts", "centres_interet", defaut=[])
        if interets:
            p = doc.add_paragraph()
            run = p.add_run("Intérêts : ")
            run.bold = True
            run = p.add_run(", ".join(interets))
        else:
            p = doc.add_paragraph()
            run = p.add_run("Intérêts : ")
            run.bold = True
            run = p.add_run("Développement web, IA & automatisation, Cybersécurité, Mangas, Sport")

        # === SAUVEGARDER ===
        entreprise = offre.get("entreprise", "entreprise")
        titre_poste = offre.get("titre", "poste")
        nom_safe = nom_fichier_safe(f"{entreprise}_{titre_poste}")

        chemin = output_dir / f"CV_{nom_safe}.docx"
        doc.save(str(chemin))
        log.info(f"CV adapté généré : {chemin.name}")
        print(f"    ✅ CV adapté : {chemin.name}")

        return chemin

    # ------------------------------------------------------------------
    # Helpers internes
    # ------------------------------------------------------------------

    @staticmethod
    def _titre_section(doc: Document, text: str):
        """Ajoute un titre de section avec ligne décorative."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 60, 90)
        # Ligne sous le titre
        pBdr = p._element.get_or_add_pPr()
        bdr = pBdr.makeelement(qn('w:pBdr'), {})
        bottom = bdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '1', qn('w:color'): '1E3C5A'
        })
        bdr.append(bottom)
        pBdr.append(bdr)

    @staticmethod
    def _ajouter_separateur(doc: Document):
        """Ajoute une ligne séparatrice horizontale."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        pBdr = p._element.get_or_add_pPr()
        bdr = pBdr.makeelement(qn('w:pBdr'), {})
        bottom = bdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '6',
            qn('w:space'): '1', qn('w:color'): '1E3C5A'
        })
        bdr.append(bottom)
        pBdr.append(bdr)

    @staticmethod
    def _remonter_projet_star(projets: list[dict], titre_star: str):
        """
        Remonte le projet le plus pertinent en première position.
        
        Matching amélioré : on compare par similarité de titre
        au lieu de chercher les 3 premiers mots (qui matchaient 'de', 'le', etc.).
        
        Stratégie : on normalise les deux titres et on cherche le meilleur
        ratio de mots significatifs en commun.
        """
        # Mots vides à ignorer
        MOTS_VIDES = {"le", "la", "les", "de", "du", "des", "un", "une",
                      "en", "et", "à", "a", "l", "d", "au", "aux"}

        def mots_significatifs(texte: str) -> set[str]:
            return {
                m.lower() for m in re.split(r"\W+", texte)
                if m and m.lower() not in MOTS_VIDES and len(m) > 1
            }

        mots_star = mots_significatifs(titre_star)
        if not mots_star:
            return

        meilleur_score = 0
        meilleur_index = -1

        for i, proj in enumerate(projets):
            mots_proj = mots_significatifs(proj["titre"])
            if not mots_proj:
                continue
            # Nombre de mots en commun / total de mots dans la requête star
            score = len(mots_star & mots_proj) / len(mots_star)
            if score > meilleur_score:
                meilleur_score = score
                meilleur_index = i

        # Remonter seulement si matching significatif (> 40%)
        if meilleur_index > 0 and meilleur_score > 0.4:
            projets.insert(0, projets.pop(meilleur_index))
            log.debug(f"Projet star remonté : '{projets[0]['titre']}' (score match: {meilleur_score:.0%})")
