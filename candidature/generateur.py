"""
generateur.py — Génération de lettres de motivation (v4)
=========================================================

V4 — Corrigé : récursion infinie RateLimitError, infos perso externalisées,
     timeout API, accents dans noms de fichiers.
"""

import json
import time
import os
import re
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import Optional

try:
    import anthropic
except ImportError:
    raise ImportError("pip install anthropic")

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

from logger import get_logger

log = get_logger("candidature.lettre")

PROFIL_PATH = Path(__file__).parent.parent / "qualification" / "profil.yaml"
OUTPUT_DIR = Path(__file__).parent / "lettres"

MAX_RETRIES = 3
RETRY_DELAYS = [10, 30, 60]


def charger_profil(chemin: Path = PROFIL_PATH) -> str:
    if not chemin.exists():
        raise FileNotFoundError(f"Profil introuvable : {chemin}")
    return chemin.read_text(encoding="utf-8")


def charger_profil_dict(chemin: Path = PROFIL_PATH) -> dict:
    """Charge le profil en dict pour extraire les infos personnelles."""
    texte = charger_profil(chemin)
    if HAS_YAML:
        return yaml.safe_load(texte) or {}
    profil = {}
    for ligne in texte.splitlines():
        if ":" in ligne and not ligne.strip().startswith("#"):
            cle, _, valeur = ligne.partition(":")
            profil[cle.strip()] = valeur.strip()
    return profil


def nom_fichier_safe(texte: str, max_len: int = 60) -> str:
    """Transforme un texte en nom de fichier ASCII safe."""
    texte = unicodedata.normalize("NFKD", texte)
    texte = texte.encode("ascii", "ignore").decode("ascii")
    texte = texte.replace(" ", "_").replace("/", "-")
    texte = re.sub(r"[^a-zA-Z0-9_-]", "", texte)
    return texte[:max_len]


def construire_prompt_lettre(profil: str, offre: dict) -> str:
    """Construit le prompt pour générer UNE lettre de motivation."""
    points_forts = offre.get("points_forts", [])
    points_faibles = offre.get("points_faibles", [])
    conseil = offre.get("conseil", "")
    score = offre.get("score", "N/A")

    return f"""Tu es un expert en rédaction de lettres de motivation pour des alternances en informatique en France.

PROFIL DU CANDIDAT :
{profil}

OFFRE VISÉE :
- Titre : {offre.get('titre', 'Non précisé')}
- Entreprise : {offre.get('entreprise', 'Non précisé')}
- Lieu : {offre.get('lieu', 'Non précisé')}
- Type de contrat : {offre.get('type_contrat', 'Alternance')}
- Description : {offre.get('description', 'Non précisé')}

ANALYSE DU MATCHING (score {score}/100) :
- Points forts : {', '.join(points_forts) if points_forts else 'Non disponible'}
- Points faibles : {', '.join(points_faibles) if points_faibles else 'Non disponible'}
- Conseil : {conseil}

CONSIGNES :
Rédige une lettre de motivation professionnelle et personnalisée.

Structure OBLIGATOIRE :
1. ACCROCHE (2-3 phrases) : Mentionne le poste exact et l'entreprise. Montre que tu connais l'entreprise.
2. PARCOURS (3-4 phrases) : Mets en avant les compétences et expériences du candidat qui matchent avec l'offre. Mentionne des projets concrets.
3. MOTIVATION (3-4 phrases) : Explique pourquoi cette alternance en particulier, pas juste "je suis motivé". Fais le lien entre les intérêts du candidat et les missions du poste.
4. APPORT MUTUEL (2-3 phrases) : Ce que le candidat apporte à l'entreprise ET ce qu'il espère y apprendre.
5. CONCLUSION (1-2 phrases) : Invitation à un entretien, formule de politesse courte.

RÈGLES :
- Ton : professionnel mais naturel (pas de phrases robotiques)
- INTERDIT : "je suis motivé et dynamique", "votre entreprise m'attire", formulations vagues
- Mentionne au moins 1 projet concret du candidat
- Adapte le vocabulaire technique au domaine de l'offre
- Maximum 350 mots
- N'inclus PAS l'en-tête (nom, adresse, date) — on l'ajoute séparément
- Commence directement par "Madame, Monsieur,"

Réponds UNIQUEMENT avec le texte de la lettre, sans aucun commentaire avant ou après."""


class Generateur:
    """Génère des lettres de motivation via l'API Claude."""

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-20250514", timeout: float = 30.0):
        self.client = anthropic.Anthropic(api_key=api_key, timeout=timeout)
        self.model = model
        self.profil = charger_profil()
        self.profil_dict = charger_profil_dict()
        OUTPUT_DIR.mkdir(exist_ok=True)
        log.info(f"Générateur initialisé (modèle: {model})")

    # --- Infos personnelles depuis profil.yaml ---
    def _info(self, *cles: str, defaut: str = "") -> str:
        for cle in cles:
            val = self.profil_dict.get(cle)
            if val:
                return str(val)
        return defaut

    @property
    def nom(self) -> str:
        return self._info("nom", defaut="Candidat")

    @property
    def email(self) -> str:
        return self._info("email", "mail", defaut="")

    @property
    def telephone(self) -> str:
        return self._info("telephone", "tel", "téléphone", defaut="")

    @property
    def localisation(self) -> str:
        return self._info("localisation", "lieu", "region", defaut="")

    def generer_lettre(self, offre: dict, _tentative: int = 0) -> Optional[str]:
        """Génère le texte d'une lettre de motivation."""
        prompt = construire_prompt_lettre(self.profil, offre)

        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}],
            )
            return response.content[0].text.strip()

        except anthropic.RateLimitError:
            if _tentative >= MAX_RETRIES:
                log.error(f"Rate limit persistant après {MAX_RETRIES} tentatives, abandon")
                return None
            delai = RETRY_DELAYS[min(_tentative, len(RETRY_DELAYS) - 1)]
            log.warning(f"Rate limit — pause {delai}s (tentative {_tentative + 1}/{MAX_RETRIES})")
            time.sleep(delai)
            return self.generer_lettre(offre, _tentative=_tentative + 1)

        except anthropic.APITimeoutError:
            log.error("Timeout API pour génération de lettre")
            return None

        except Exception as e:
            log.error(f"Erreur API : {e}")
            return None

    def sauvegarder_docx(self, texte: str, offre: dict, chemin: Path):
        """Sauvegarde la lettre en format Word (.docx)."""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # En-tête depuis profil.yaml
        p = doc.add_paragraph()
        run = p.add_run(self.nom)
        run.bold = True
        run.font.size = Pt(12)
        if self.localisation:
            doc.add_paragraph(self.localisation)
        if self.email:
            doc.add_paragraph(self.email)
        if self.telephone:
            doc.add_paragraph(self.telephone)

        doc.add_paragraph()
        date_str = datetime.now().strftime("%d/%m/%Y")
        p_date = doc.add_paragraph(f"Le {date_str}")
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()
        entreprise = offre.get("entreprise", "l'entreprise")
        p_dest = doc.add_paragraph()
        run = p_dest.add_run(entreprise)
        run.bold = True

        doc.add_paragraph()
        titre_poste = offre.get("titre", "le poste")
        p_obj = doc.add_paragraph()
        run = p_obj.add_run(f"Objet : Candidature — {titre_poste}")
        run.bold = True

        doc.add_paragraph()
        for paragraphe in texte.split("\n\n"):
            paragraphe = paragraphe.strip()
            if paragraphe:
                p = doc.add_paragraph(paragraphe)
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(str(chemin))

    def sauvegarder_pdf(self, texte: str, offre: dict, chemin: Path):
        """Sauvegarde la lettre en format PDF."""
        from fpdf import FPDF
        import platform

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=25)

        font_loaded = False
        if platform.system() == "Windows":
            font_paths = [
                ("C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/calibrib.ttf"),
                ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
            ]
            for regular, bold in font_paths:
                if os.path.exists(regular):
                    pdf.add_font("CustomFont", "", regular, uni=True)
                    pdf.add_font("CustomFont", "B", bold, uni=True)
                    font_loaded = True
                    break
        else:
            linux_fonts = [
                ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                 "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            ]
            for regular, bold in linux_fonts:
                if os.path.exists(regular):
                    pdf.add_font("CustomFont", "", regular, uni=True)
                    pdf.add_font("CustomFont", "B", bold, uni=True)
                    font_loaded = True
                    break

        font_name = "CustomFont" if font_loaded else "Helvetica"

        # En-tête depuis profil.yaml
        pdf.set_font(font_name, "B", 12)
        pdf.cell(0, 7, self.nom, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 10)
        if self.localisation:
            pdf.cell(0, 6, self.localisation, new_x="LMARGIN", new_y="NEXT")
        if self.email:
            pdf.cell(0, 6, self.email, new_x="LMARGIN", new_y="NEXT")
        if self.telephone:
            pdf.cell(0, 6, self.telephone, new_x="LMARGIN", new_y="NEXT")

        pdf.ln(8)
        date_str = datetime.now().strftime("%d/%m/%Y")
        pdf.cell(0, 6, f"Le {date_str}", align="R", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(6)
        entreprise = offre.get("entreprise", "l'entreprise")
        pdf.set_font(font_name, "B", 10)
        pdf.cell(0, 6, entreprise, new_x="LMARGIN", new_y="NEXT")

        pdf.ln(6)
        titre_poste = offre.get("titre", "le poste")
        pdf.set_font(font_name, "B", 10)
        pdf.cell(0, 6, f"Objet : Candidature - {titre_poste}", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(8)
        pdf.set_font(font_name, "", 10)
        for paragraphe in texte.split("\n\n"):
            paragraphe = paragraphe.strip()
            if paragraphe:
                pdf.multi_cell(0, 6, paragraphe)
                pdf.ln(4)

        pdf.output(str(chemin))

    def generer_pour_offre(self, offre: dict) -> Optional[dict]:
        """Pipeline complet : génère texte → docx → pdf."""
        titre = offre.get("titre", "offre")
        entreprise = offre.get("entreprise", "entreprise")
        score = offre.get("score", 0)

        log.info(f"✍️ Génération pour : {titre} @ {entreprise} (score: {score})")

        texte = self.generer_lettre(offre)
        if not texte:
            return None

        nom_safe = nom_fichier_safe(f"{entreprise}_{titre}")

        chemin_docx = OUTPUT_DIR / f"LM_{nom_safe}.docx"
        try:
            self.sauvegarder_docx(texte, offre, chemin_docx)
            log.info(f"📄 Word : {chemin_docx.name}")
        except Exception as e:
            log.warning(f"Erreur DOCX : {e}")
            chemin_docx = None

        chemin_pdf = OUTPUT_DIR / f"LM_{nom_safe}.pdf"
        try:
            self.sauvegarder_pdf(texte, offre, chemin_pdf)
            log.info(f"📕 PDF : {chemin_pdf.name}")
        except Exception as e:
            log.warning(f"Erreur PDF : {e}")
            chemin_pdf = None

        chemin_txt = OUTPUT_DIR / f"LM_{nom_safe}.txt"
        chemin_txt.write_text(texte, encoding="utf-8")

        return {
            "offre_id": offre.get("id"),
            "entreprise": entreprise, "titre": titre, "score": score,
            "docx": str(chemin_docx) if chemin_docx else None,
            "pdf": str(chemin_pdf) if chemin_pdf else None,
            "txt": str(chemin_txt),
        }
